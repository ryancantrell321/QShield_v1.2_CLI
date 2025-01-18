import logging
from logging.handlers import RotatingFileHandler
import os
from docx import Document
from docx.shared import Pt
from datetime import datetime

try:
    # Constructing the log directory path near the main.py directory
    documents_directory = os.path.expanduser("~/Documents")
    log_directory = os.path.join(documents_directory, "qShield Event Logs")
    os.makedirs(log_directory, exist_ok=True)

    # Log file paths
    LOG_FILE = os.path.join(log_directory, "events.log")
    # DOCX_FILE = os.path.join(log_directory, "events.docx")

    # Setting up the logger
    logger = logging.getLogger("QShield Logger")
    logger.setLevel(logging.DEBUG)
    # Create a rotating file handler that logs to a file with max size of 1MB and keeps 3 backup files
    file_handler = RotatingFileHandler(LOG_FILE, maxBytes=1*1024*1024, backupCount=3)
    file_handler.setLevel(logging.DEBUG)

    # Create a console handler to also output logs to the console
    console_handler = logging.StreamHandler()
    console_handler.setLevel(logging.DEBUG)

    # Create a formatter and set it for the handlers
    formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    file_handler.setFormatter(formatter)
    console_handler.setFormatter(formatter)

    # Add the handlers to the logger
    logger.addHandler(file_handler)
    # logger.addHandler(console_handler)

    # # Create/open the .docx file
    # if os.path.exists(DOCX_FILE):
    #     doc = Document(DOCX_FILE)
    # else:
    #     doc = Document()
    #     doc.add_heading("Pandora QShield Event Logs", 0)
    #
    # def log_to_docx(message, level="INFO"):
    #     # Add a new paragraph with the log message
    #     p = doc.add_paragraph()
    #     run = p.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')} - {level} - {message}")
    #     run.font.size = Pt(13)
    #     if level == "ERROR":
    #         run.font.bold = True
    #     doc.save(DOCX_FILE)

    def log_message(level, message):
        if level == "DEBUG":
            logger.debug(message)
        elif level == "INFO":
            logger.info(message)
        elif level == "WARNING":
            logger.warning(message)
        elif level == "ERROR":
            logger.error(message)
        elif level == "CRITICAL":
            logger.critical(message)
        # log_to_docx(message, level)

    # Example usage
    if __name__ == "__main__":
        log_message("INFO", "This is an info message")
        log_message("ERROR", "This is an error message")

except Exception as e:
    print(f"Error: {e}")
